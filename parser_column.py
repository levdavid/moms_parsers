# -*- coding: utf-8 -*- 
from docx.api import Document


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

# Load the first table from your document. In your example file,
# there is only one table, so I just grab the first one.
document = Document('10.docx')

story = document.tables[1]

# Data will be a list of rows represented as dictionaries
# containing each row's data.
data = []

keys = None
for i, row in enumerate(story.rows):
    cur_cell = row.cells[3]
    for paragraph in cur_cell.paragraphs:
        for run in paragraph.runs:
            if not run.italic:
                paragraph.clear()
                paragraph.style = None
                break
document.save( '../' + unicode('серия_10', 'utf-8')+ '.docx')
