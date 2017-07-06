# -*- coding: utf-8 -*- 
from docx.api import Document

# Load the first table from your document. In your example file,
# there is only one table, so I just grab the first one.
document = Document('13.docx')

heroes_table = document.tables[0]

heroes = []

keys = None
for i, row in enumerate(heroes_table.rows):
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if run.bold:
                    name = run.text.split(' ')[0]
                    if len(name) > 2 and name[len(name)-1]!=':':
                        heroes.append(u''.join(name))

for hero in  heroes :
    print hero

story = document.tables[1]

# Data will be a list of rows represented as dictionaries
# containing each row's data.
data = []

keys = None
for i, row in enumerate(story.rows):
    text = (cell.text for cell in row.cells)

    # Establish the mapping based on the first row
    # headers; these will become the keys of our dictionary
    if i == 0:
        keys = tuple(text)
        continue

    # Construct a dictionary for this row, mapping
    # keys to values for this row
    row_data = tuple(text)
    data.append(row_data)

for hero in heroes:
    hero_doc = Document()

    hero_doc.add_heading(hero, level=1)

    table = hero_doc.add_table(rows=1, cols=4)
    table.style = 'TableGrid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = unicode('Номер сцены', 'utf-8')
    hdr_cells[0].width = 1097280    
    hdr_cells[1].text = unicode('Описание', 'utf-8')
    hdr_cells[1].width = 4846320    
    hdr_cells[2].text = unicode('Персонажи', 'utf-8')
    hdr_cells[2].width = 2423160
    hdr_cells[3].text = unicode('Примечания', 'utf-8')
    hdr_cells[3].width = 20583680
    for row in data:
        if hero in u''.join(row[4]):
            row_cells = table.add_row().cells
            row_cells[0].width = 1097280    
            row_cells[1].width = 4846320    
            row_cells[2].width = 2423160
            row_cells[3].width = 20583680

            row_cells[0].text = u''.join(row[0])
            row_cells[1].text = u''.join(row[1])
            row_cells[2].text = u''.join(row[4])



    hero_doc.save( unicode('серия_13', 'utf-8')+ "_" + hero +'.docx')
